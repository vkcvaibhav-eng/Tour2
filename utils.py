import os
import google.generativeai as genai
import pandas as pd
import json
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. GEMINI API CALL ---
def call_gemini_extraction(api_key, file_paths, prompt):
    """
    Uploads files to Gemini and extracts data based on the prompt.
    """
    if not api_key:
        return '{"error": "API Key missing"}'

    genai.configure(api_key=api_key)
    
    # Upload files to Gemini
    uploaded_files = []
    for path in file_paths:
        try:
            # Uploading file
            sample_file = genai.upload_file(path=path, display_name=os.path.basename(path))
            uploaded_files.append(sample_file)
        except Exception as e:
            return f'{{"error": "File upload failed: {str(e)}"}}'

    # Configure Model
    model = genai.GenerativeModel(model_name="gemini-3-flash-preview")

    # Generate Content
    try:
        response = model.generate_content([prompt, *uploaded_files])
        return response.text
    except Exception as e:
        return f'{{"error": "AI Generation failed: {str(e)}"}}'

# --- 2. JSON CLEANER ---
def clean_and_parse_json(raw_text):
    """
    Cleans Markdown formatting (```json ... ```) and parses JSON.
    """
    try:
        # Remove markdown code blocks
        clean_text = raw_text.replace("```json", "").replace("```", "").strip()
        return json.loads(clean_text)
    except Exception:
        # Fallback: Try finding the first { and last }
        try:
            start = raw_text.find("{")
            end = raw_text.rfind("}") + 1
            if start != -1 and end != -1:
                return json.loads(raw_text[start:end])
            return {"error": "Could not parse JSON", "raw_text": raw_text}
        except:
            return {"error": "Critical JSON parsing error"}

# --- 3. WORD DOCUMENT GENERATOR (EXPORT) ---
def create_complex_claim_form(diary_df, ta_df, da_df):
    """
    Generates the Word document.
    """
    doc = Document()
    
    # Title
    head = doc.add_heading('TA/DA CLAIM FORM', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- SECTION 1: TOUR DIARY ---
    doc.add_heading('1. Tour Diary', level=1)
    
    if not diary_df.empty:
        # Create Table
        table = doc.add_table(rows=1, cols=len(diary_df.columns))
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(diary_df.columns):
            hdr_cells[i].text = str(col_name)
        
        # Rows
        for _, row in diary_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
    else:
        doc.add_paragraph("No Tour Diary Data")

    # --- SECTION 2: TA CALCULATION ---
    doc.add_heading('2. Transport Allowance (TA)', level=1)
    
    if not ta_df.empty:
        # Filter relevant columns for export
        cols_to_print = ["Departure_Place", "Arrival_Place", "Mode_of_Travel", "KM", "Total_Amount"]
        # Ensure columns exist
        actual_cols = [c for c in cols_to_print if c in ta_df.columns]
        
        table = doc.add_table(rows=1, cols=len(actual_cols))
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(actual_cols):
            hdr_cells[i].text = col
            
        for _, row in ta_df.iterrows():
            row_cells = table.add_row().cells
            for i, col in enumerate(actual_cols):
                val = row[col]
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.2f}"
                else:
                    row_cells[i].text = str(val)
                    
        # Total TA
        total_ta = ta_df["Total_Amount"].sum() if "Total_Amount" in ta_df else 0
        p = doc.add_paragraph()
        p.add_run(f"Total Transport Allowance: ₹ {total_ta:,.2f}").bold = True

    # --- SECTION 3: DA CALCULATION ---
    doc.add_heading('3. Daily Allowance (DA)', level=1)
    
    if not da_df.empty:
        table = doc.add_table(rows=1, cols=len(da_df.columns))
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(da_df.columns):
            hdr_cells[i].text = str(col)
            
        for _, row in da_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

        # Total DA
        total_da = da_df["Total_DA"].sum() if "Total_DA" in da_df.columns else 0
        p = doc.add_paragraph()
        p.add_run(f"Total Daily Allowance: ₹ {total_da:,.2f}").bold = True

    # --- GRAND TOTAL ---
    doc.add_paragraph("\n")
    grand_total = (ta_df["Total_Amount"].sum() if "Total_Amount" in ta_df else 0) + \
                  (da_df["Total_DA"].sum() if "Total_DA" in da_df.columns else 0)
                  
    final_p = doc.add_paragraph()
    runner = final_p.add_run(f"GRAND TOTAL CLAIM: ₹ {grand_total:,.2f}")
    runner.bold = True
    runner.font.size = Pt(14)
    final_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc

