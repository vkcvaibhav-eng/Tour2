import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Cm, Mm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

# --- PAGE CONFIGURATION ---
st.set_page_config(layout="wide", page_title="Step 5: Final Gujarati Export")
st.title("🇮🇳 Step 5: Final Gujarati Export (Split Layout)")
st.markdown("---")

# --- 1. SMART INPUTS FOR HEADER DATA ---
st.sidebar.header("📝 Header Details")

# Defaults based on your prompt
emp_name = st.sidebar.text_input("Employee Name", "વી. કે. ચૌધરી")
designation = st.sidebar.text_input("Designation", "સીનીયર એક્રોલોજીસ્ટ")
basic_pay = st.sidebar.text_input("Basic Pay", "₹ 67,700") 
pay_level = st.sidebar.text_input("Pay Level", "Level-11")
pay_scale = st.sidebar.text_input("Pay Scale", "67700 - 208700")
office_name = st.sidebar.text_input("Office", "કીટકશાસ્ત્ર વિભાગ")
work_place = st.sidebar.text_input("Work Place", "કીટકશાસ્ત્ર વિભાગ, ન.મ.કૃ.મ.,ન.કૃ.યુ., નવસારી")
voucher_month = st.sidebar.text_input("Bill Month", "ફેબ્રુઆરી - ૨૦૨૬")

# --- DATA CONNECTION ---
if 'final_18_col_df' in st.session_state:
    df_ta = st.session_state['final_18_col_df'].copy()
    # Add Column 19 (Remarks) as empty if missing
    if len(df_ta.columns) == 18:
        df_ta["19. Remarks"] = ""
    st.success(f"✅ Connected to Final Table: {len(df_ta)} rows.")
else:
    df_ta = pd.DataFrame()
    st.warning("⚠️ No data found. Please complete previous steps.")

# --- HELPER: GUJARATI NUMBER TO WORDS ---
def num_to_gujarati_words(num):
    try:
        num = int(num)
    except:
        return ""
    
    ones = ["", "એક", "બે", "ત્રણ", "ચાર", "પાંચ", "છ", "સાત", "આઠ", "નવ"]
    teens = ["દસ", "અગિયાર", "બાર", "તેર", "ચૌદ", "પંદર", "સોળ", "સત્તર", "અઢાર", "ઓગણીસ"]
    tens = ["", "", "વીસ", "ત્રીસ", "ચાલીસ", "પચાસ", "સાઈઠ", "સિત્તેર", "એંસી", "નેવું"]
    
    def convert_upto_99(n):
        if n < 10: return ones[n]
        if 10 <= n < 20: return teens[n-10]
        return tens[n // 10] + (" " + ones[n % 10] if n % 10 != 0 else "")

    parts = []
    if num >= 100000:
        parts.append(convert_upto_99(num // 100000) + " લાખ")
        num %= 100000
    if num >= 1000:
        parts.append(convert_upto_99(num // 1000) + " હજાર")
        num %= 1000
    if num >= 100:
        parts.append(convert_upto_99(num // 100) + " સો")
        num %= 100
    if num > 0:
        parts.append(convert_upto_99(num))
        
    return " ".join(parts).strip()

# --- DOCX FORMATTING HELPERS ---
def format_cell(cell, text, font_size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    p.alignment = align
    for run in p.runs:
        run._element.getparent().remove(run._element)
    run = p.add_run(str(text))
    run.font.name = 'Arial Unicode MS'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_row_height(row, height_cm):
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

# --- MAIN GENERATOR FUNCTION ---
def create_split_layout_doc(ta_data):
    doc = Document()
    
    # 1. PRE-CALCULATE TOTALS (Required for Page 1 & 4)
    total_claim = pd.to_numeric(ta_data.iloc[:, 16].astype(str).str.replace(r'[^\d.]', '', regex=True), errors='coerce').sum()
    total_words = num_to_gujarati_words(int(total_claim))
    
    # GLOBAL STYLES
    style = doc.styles['Normal']
    style.font.name = 'Arial Unicode MS'
    style.font.size = Pt(11)

    # ================= PAGE 1: HISABI PATRAK =================
    # Use existing first section instead of adding new one
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(20)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    # Top Right: Hisabi Patrak
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("હિસાબી પત્રક નંબર ____________")
    
    # Center Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("નવસારી કૃષિ વિશ્વવિધાલય\n")
    run.bold = True
    run.font.size = Pt(16)
    run = p.add_run("મુસાફરી ભથ્થા બીલ")
    run.bold = True
    run.font.size = Pt(14)

    # Info Table (Bill No vs Voucher No)
    table = doc.add_table(rows=1, cols=2)
    table.width = Mm(175)
    table.autofit = False
    
    # Left Column
    cell_l = table.cell(0, 0)
    cell_l.width = Mm(80)
    p = cell_l.paragraphs[0]
    p.add_run("બીલ નંબર :\nતારીખ        :")
    p.runs[0].bold = True

    # Right Column
    cell_r = table.cell(0, 1)
    cell_r.width = Mm(95)
    p = cell_r.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("વાઉચર નં. ____________________\n")
    p.add_run("તારીખ ____________________\n")
    p.add_run("યુનિટ નંબર : __________________\n")
    p.add_run("કોડ નંબર : ____________________")

    doc.add_paragraph() # Spacer

    # Main Body Text
    p = doc.add_paragraph(f"આચાર્ય અને ડીનશ્રી, નં. મ. કૃષિ મહાવિદ્યાલય, નકૃયું, નવસારી ની કચેરીનું માહે: {voucher_month} નું")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("મુસાફરી ભથ્થા બિલ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True
    p.runs[0].font.size = Pt(14)

    p = doc.add_paragraph("યુનિટ/સબયુનિટ : આચાર્ય અને ડીનશ્રી, ન. મ. કૃષિ મહાવિદ્યાલય, નકૃયું, નવસારી")
    p = doc.add_paragraph("ખર્ચ માટેનું બજેટ સદર :- ____________________")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("યોજનાનું નામ :- ____________________")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() # Spacer

    # Claim Amount
    p = doc.add_paragraph()
    p.add_run("આથી રૂ।. ")
    r = p.add_run(f" {total_claim:.0f} ")
    r.bold = True
    r.underline = True
    p.add_run(" નો દાવો મંજુર કરી ગ્રાહય રાખવામાં આવે છે.")

    # Boxed Amount Area (Simulated with borders)
    tbl_amt = doc.add_table(rows=1, cols=1)
    tbl_amt.style = 'Table Grid'
    cell = tbl_amt.cell(0,0)
    p = cell.paragraphs[0]
    p.add_run("આ બીલમાં જણાવેલ રૂા  ")
    p.add_run(f"{total_claim:.0f}").bold = True
    p.add_run("  ( અંકે રૂપિયા ")
    p.add_run(f"{total_words} પુરા").bold = True
    p.add_run(" પૈસા )\n\n")
    p.add_run("મંજુર કરવામાં આવે છે. અને તે રોકડા / ચેક નં. ______________ તા. ___________ થી ચુકવવામાં આવે છે.")

    doc.add_paragraph() # Spacer

    # Signatures
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.width = Mm(175)
    
    l_cell = sig_table.cell(0, 0)
    l_cell.text = "સ્થળ :    નવસારી\nતારીખ :"
    
    r_cell = sig_table.cell(0, 1)
    p = r_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("____________________\n")
    p.add_run("બીલ મંજુર કરનાર અધિકારીની\nસહી અને હોદ્દો")

    doc.add_paragraph() # Spacer

    # Budget Table
    b_table = doc.add_table(rows=4, cols=3)
    b_table.style = 'Table Grid'
    hdr_cells = b_table.rows[0].cells
    hdr_cells[0].text = ""
    hdr_cells[1].text = "રૂ."
    hdr_cells[2].text = "પૈસા"
    
    b_table.cell(1, 0).text = "(૧) સને ૨૦૨૪-૨૫ માટે બજેટમાં મંજુર થયેલ રકમ"
    b_table.cell(2, 0).text = "(૨) આ બીલ સાથે થયેલ કુલ ખર્ચ"
    b_table.cell(3, 0).text = "(૩) ખર્ચ માટે બાકી રહેતી રકમ"

    doc.add_paragraph() # Spacer

    # Bottom Signature
    p = doc.add_paragraph()
    p.add_run(f"રૂા ( {total_claim:.0f} ) અંકે રૂપિયા : {total_words} પુરા મંજુર કર્યા")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("નિયંત્રણ અધિકારીની સહી___________").bold = True

    doc.add_page_break()

    # ================= PAGE 2: TABLE PART 1 (Cols 1-10) =================
    
    # NEW HEADER FOR PAGE 2 (LEFT SIDE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = p.add_run(f"કાર્ય મથક : {work_place}\n")
    run.font.name = 'Arial Unicode MS'; run.font.size = Pt(11); run.bold = True
    
    run = p.add_run(f"કચેરી : {office_name}\n")
    run.font.name = 'Arial Unicode MS'; run.font.size = Pt(11); run.bold = True
    
    run = p.add_run(f"કર્મચારીનુ નામ: {emp_name}")
    run.font.name = 'Arial Unicode MS'; run.font.size = Pt(11); run.bold = True

    doc.add_paragraph() # Spacer before table

    # Columns: 1.DepPlace, 2.Date, 3.Time, 4.ArrPlace, 5.Date, 6.Time, 7.Mode, 8.Class, 9.Price, 10.Total
    table_left = doc.add_table(rows=3, cols=10)
    table_left.style = 'Table Grid'
    table_left.autofit = False
    
    # Set Column Widths (Approx to fit A4 width ~18cm)
    col_widths_L = [2.0, 1.8, 1.3, 2.0, 1.8, 1.3, 2.5, 1.5, 1.8, 2.0] # Total ~18cm
    for r in table_left.rows:
        for i, w in enumerate(col_widths_L):
            r.cells[i].width = Cm(w)

    # --- HEADER ROWS ---
    # Row 0
    r0 = table_left.rows[0].cells
    r0[0].merge(r0[2]); format_cell(r0[0], "નીકળ્યા (Departure)", bold=True)
    r0[3].merge(r0[5]); format_cell(r0[3], "આવ્યા (Arrival)", bold=True)
    
    # Vertical Merges for Mode, Class, Ticket
    for i in range(6, 10):
        table_left.cell(0, i).merge(table_left.cell(1, i))
    
    format_cell(table_left.cell(0, 6), "મુસાફરીનો\nપ્રકાર", bold=True, font_size=9)
    format_cell(table_left.cell(0, 7), "વર્ગ", bold=True, font_size=9)
    format_cell(table_left.cell(0, 8), "ભાડા દર", bold=True, font_size=9)
    format_cell(table_left.cell(0, 9), "ટિકિટ રકમ\n(Col 10)", bold=True, font_size=9)

    # Row 1 (Sub headers for Dep/Arr)
    r1 = table_left.rows[1].cells
    headers_sub = ["સ્થળ", "તારીખ", "સમય", "સ્થળ", "તારીખ", "સમય"]
    for i, txt in enumerate(headers_sub):
        format_cell(r1[i], txt, font_size=9)

    # Row 2 (Numbers)
    for i in range(10):
        format_cell(table_left.rows[2].cells[i], str(i+1), bold=True)

    # --- DATA ROWS PART 1 ---
    row_height_cm = 1.0 # STRICT HEIGHT
    
    for index, row in ta_data.iterrows():
        new_row = table_left.add_row()
        set_row_height(new_row, row_height_cm) # SYNCHRONIZATION KEY
        
        cells = new_row.cells
        # Col 1-6
        for i in range(6):
            val = str(row.iloc[i]) if pd.notnull(row.iloc[i]) else ""
            format_cell(cells[i], val, font_size=9)
        # Col 7 Mode
        format_cell(cells[6], str(row.iloc[6]), font_size=9)
        # Col 8 Class
        format_cell(cells[7], str(row.iloc[7]), font_size=9)
        # Col 9 Price
        try: pr = float(str(row.iloc[8]).replace('₹','').replace(',',''))
        except: pr = 0.0
        format_cell(cells[8], f"{pr:.0f}", font_size=9)
        # Col 10 Total
        try: tot = float(str(row.iloc[9]).replace('₹','').replace(',',''))
        except: tot = 0.0
        format_cell(cells[9], f"{tot:.0f}", font_size=9, bold=True)

    # Add Total Row
    tot_row_L = table_left.add_row()
    set_row_height(tot_row_L, row_height_cm)
    tot_row_L.cells[0].merge(tot_row_L.cells[8])
    format_cell(tot_row_L.cells[0], "સરવાળો (Page Total)", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    
    # Sum Col 10
    sum_10 = pd.to_numeric(ta_data.iloc[:, 9].astype(str).str.replace(r'[^\d.]', '', regex=True), errors='coerce').sum()
    format_cell(tot_row_L.cells[9], f"{sum_10:.0f}", bold=True)

    doc.add_page_break()

    # ================= PAGE 3: TABLE PART 2 (Cols 11-19) =================
    
    # NEW HEADER FOR PAGE 3 (RIGHT SIDE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = p.add_run(f"હોદ્દો: {designation}\n")
    run.font.name = 'Arial Unicode MS'; run.font.size = Pt(11); run.bold = True
    
    run = p.add_run(f"મૂળ પગાર: {basic_pay}\n")
    run.font.name = 'Arial Unicode MS'; run.font.size = Pt(11); run.bold = True
    
    run = p.add_run(f"Level: {pay_level}  |  Pay scale: {pay_scale}")
    run.font.name = 'Arial Unicode MS'; run.font.size = Pt(11); run.bold = True

    doc.add_paragraph() # Spacer before table

    # Columns: 11.KM, 12.Rate, 13.Total, 14.Days, 15.DA_Rate, 16.DA_Amt, 17.GrandTotal, 18.Purpose, 19.Remarks
    table_right = doc.add_table(rows=3, cols=9)
    table_right.style = 'Table Grid'
    table_right.autofit = False

    # Widths
    col_widths_R = [1.5, 1.5, 2.0, 1.5, 1.5, 2.0, 2.5, 3.5, 2.0]
    for r in table_right.rows:
        for i, w in enumerate(col_widths_R):
            r.cells[i].width = Cm(w)

    # --- HEADER ROWS ---
    # Row 0
    r0 = table_right.rows[0].cells
    r0[0].merge(r0[2]); format_cell(r0[0], "અન્ય વાહન મુસાફરી (Road)", bold=True, font_size=9)
    r0[3].merge(r0[5]); format_cell(r0[3], "દૈનિક ભથ્થું (DA)", bold=True, font_size=9)
    
    # Vertical Merges
    for i in range(6, 9):
        table_right.cell(0, i).merge(table_right.cell(1, i))
        
    format_cell(table_right.cell(0, 6), "કુલ રકમ\n(૧૦+૧૩+૧૬)", bold=True, font_size=9)
    format_cell(table_right.cell(0, 7), "મુસાફરીનું કારણ", bold=True, font_size=9)
    format_cell(table_right.cell(0, 8), "વિશેષ નોંધ", bold=True, font_size=9)

    # Row 1 (Sub headers)
    r1 = table_right.rows[1].cells
    sub_headers_R = ["કિ.મી.", "દર", "રકમ", "દિવસ", "દર", "રકમ"]
    for i, txt in enumerate(sub_headers_R):
        format_cell(r1[i], txt, font_size=9)

    # Row 2 (Numbers 11-19)
    for i in range(9):
        format_cell(table_right.rows[2].cells[i], str(i+11), bold=True)

    # --- DATA ROWS PART 2 ---
    # Iterate same data again to keep rows matched
    for index, row in ta_data.iterrows():
        new_row = table_right.add_row()
        set_row_height(new_row, row_height_cm) # CRITICAL SYNC
        
        cells = new_row.cells
        
        # Helper to clean numbers
        def clean(idx): 
            try: return float(str(row.iloc[idx]).replace('₹','').replace(',',''))
            except: return 0.0

        # Col 11-13 (Road)
        km = clean(10); rt = clean(11); r_tot = clean(12)
        format_cell(cells[0], f"{km:.1f}" if km else "-")
        format_cell(cells[1], f"{rt:.1f}" if rt else "-")
        format_cell(cells[2], f"{r_tot:.0f}" if r_tot else "-")
        
        # Col 14-16 (DA)
        d_day = row.iloc[13]; d_rate = clean(14); d_amt = clean(15)
        format_cell(cells[3], str(d_day) if str(d_day)!="0" else "-")
        format_cell(cells[4], f"{d_rate:.0f}" if d_rate else "-")
        format_cell(cells[5], f"{d_amt:.0f}" if d_amt else "-")
        
        # Col 17 (Total)
        g_tot = clean(16)
        format_cell(cells[6], f"{g_tot:.0f}", bold=True)
        
        # Col 18 Purpose
        format_cell(cells[7], str(row.iloc[17]), font_size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
        
        # Col 19 Remarks
        format_cell(cells[8], "", font_size=8)

    # Total Row
    tot_row_R = table_right.add_row()
    set_row_height(tot_row_R, row_height_cm)
    tot_row_R.cells[0].merge(tot_row_R.cells[5])
    format_cell(tot_row_R.cells[0], "કુલ સરવાળો (Grand Total)", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    
    # Grand Total Calculation
    grand_sum = pd.to_numeric(ta_data.iloc[:, 16].astype(str).str.replace(r'[^\d.]', '', regex=True), errors='coerce').sum()
    format_cell(tot_row_R.cells[6], f"₹ {grand_sum:.0f}", bold=True)
    
    doc.add_page_break()

    # ================= PAGE 4: NOTES & CERTIFICATES =================
    
    p = doc.add_paragraph("નોંધ :-")
    p.runs[0].bold = True
    
    notes = [
        "કોલમ નં. ૭ માં મુસાફરી પ્રકાર રેલ્વે/એસ.ટી./હવાઈ/સ્ટીમર/ભાડાનું યુનિવર્સિટી કે સરકારી કે પોતાનું વાહન ઈત્યાદી મારફત કરેલ મુસાફરીની સ્પષ્ટ નોંધ આપવી.",
        "કોલમ નં. ૧૧ થી ૧૩ માઈલેજ મેળવતા અધિકારીઓ કે સભ્યોએ ભરવી.",
        "કોલમ નં. ૧૬ માં માત્ર દૈનિક ભથ્થાની રકમ લેવી જેથી કોલમ (૧૪ X ૧૫= ૧૬) થઈ રહેવું જોઈએ.",
        "જયારે મુસાફરી ભથ્થા બીલમાં શરૂઆતમાં મુસાફરીને બદલે 'હોલ્ટ' દર્શાવવામાં આવેલ હોય તેવા કિસ્સામાં 'હોલ્ટ' ની શરૂઆત થયાની તારીખ કો.નં. ૧૯ માં દર્શાવવી."
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Number')

    doc.add_paragraph()
    p = doc.add_paragraph("યુનિવર્સિટી કર્મચારીએ આપવાનું પ્રમાણપત્ર")
    p.runs[0].bold = True
    
    certs = [
        "આથી પ્રમાણપત્ર આપવામાં આવે છે કે, આ બીલમાં આકારેલ રકમ બીજા કોઈ બીલમાં આકારેલ નથી.",
        "આથી પ્રમાણીત કરવામાં આવે છે કે સદર મુસાફરી ભથ્થ બીલમાં દર્શાવેલ હકીકત સાચી છે અને સદર મુસાફરી ભથ્થા બીલમાં કરવામાં આવેલ દાવો વખતો વખત સુધારવામાં આવેલ ગુજરાત કૃષિ વિશ્વવિધાલય મુસાફરી ભથ્થા નિયમોના ૬૫-૧ ના અનુમાનોને આધારે સાચો છે.",
        "આથી પ્રમાણપત્ર આપવામાં આવે છે કે બીલમાં દર્શાવેલ પ્રવાસ માટે મેં આ અગાઉ પેશગી લીધેલ નથી. / મે પેશગી લીધેલ છે જેની વિગત નીચે દર્શાવેલ છે.  (લાગુ ન પડતી બાબત છેકી નાખવી અને તે બદલ સહી કરવી.)",
        "આ બીલમાં જણાવેલ યુનિવર્સિટી સિવાયની અન્ય સંસ્થાની કામગીરીના પ્રવાસ માટે જે તે સંસ્થા તરફથી પ્રવાસ ભથ્થના નાણાં મને મળેલ નથી અને જે તે સંસ્થા તરફથી નાણાં મળશે તો તે રકમ થુનિવર્સિટી ફંડમાં જમા કરાવવામાં આવશે.",
        "આથી પ્રમાણપત્ર આપવામાં આવે છે કે, પ્રવાસ ડાયરીમાં દર્શાવવામાં આવેલ સ્થળ, તારીખ, સમય, કિલોમીટર કચેરીના વાહન લોગબુક મુજબ આકારવામાં આવેલ છે."
    ]
    for cert in certs:
        doc.add_paragraph(cert, style='List Number')

    # Employee Signature
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"({emp_name})\n").bold = True
    p.add_run(f"{designation}\nકર્મચારીની સહી નામ અને હોદ્દો")

    doc.add_paragraph("_" * 50) # Horizontal Line simulation

    # Officer Cert
    p = doc.add_paragraph("યુનિવર્સિટી અધિકારીઓ અને અન્ય સભ્યોએ આપવાનું પ્રમાણપત્ર")
    p.runs[0].bold = True
    doc.add_paragraph("આથી પ્રમાણિત કરવામાં આવે છે કે સદર બીલમાં કરેલ મુસાફરી ભથ્થાનો દાવો આ અંગેના નિયમોની જોગવાઈઓના આધારે ખરો અને યોગ્ય છે.")

    # Officer Signature
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("_______________________\n")
    p.add_run("પ્રાધ્યાપક અને વડા\nકિટકશાત્ર વિભાગ\nનં. મ. કૃષિ મહાવિદ્યાલય\nનકૃયું, નવસારી").bold = True

    doc.add_paragraph()

    # --- FINAL SPLIT TABLE (CALCULATION & RECEIPT) ---
    p = doc.add_paragraph("કર્મચારી / અધિકારી / સભ્યશ્રીએ નીચેની વિગત ભરવી.")
    p.runs[0].bold = True

    final_table = doc.add_table(rows=1, cols=2)
    final_table.style = 'Table Grid'
    final_table.allow_autofit = False
    
    # Left Cell (Calculation)
    c1 = final_table.cell(0, 0)
    c1.width = Mm(90)
    p = c1.paragraphs[0]
    p.add_run(f"બીલની કુલ રકમ\t= {total_claim:.0f}\n")
    p.add_run("બાદ બીલની પેશગીની રકમ\t=\n")
    p.add_run(f"ચૂકવવા પાત્ર ચોખ્ખી રકમ\t= {total_claim:.0f}\n\n")
    p.add_run("પેશગીના નાણાં મળ્યાની તા.\n")
    p.add_run("પેશગી કયા ઝોન /યુનિટમાંથી\t નીલ\nઉપાડવામાં આવી.\n\n")
    p.add_run("પેશગી ઉપાડવાના વાઉચર નંબર ______ તારીખ _____")

    # Right Cell (Receipt)
    c2 = final_table.cell(0, 1)
    c2.width = Mm(90)
    p = c2.paragraphs[0]
    p.add_run(f"બીલની રકમ રૂ. {total_claim:.0f}\n").bold = True
    p.add_run(f"અંકે રૂપિયા {total_words} પુરા\n").bold = True
    p.add_run("મને મળ્યા છે.\n\n")
    p.add_run("સ્થળ : નવસારી\n")
    p.add_run("તારીખ :\n\n\n")
    p.add_run(f"                                ({emp_name})\n").bold = True
    p.add_run(f"                                {designation}")

    return doc

# --- UI PREVIEW ---

# Calculate totals for preview
total_claim_val = 0.0
total_words_str = "શૂન્ય"
if not df_ta.empty:
    total_claim_val = pd.to_numeric(df_ta.iloc[:, 16], errors='coerce').sum()
    total_words_str = num_to_gujarati_words(int(total_claim_val))

# HTML PREVIEW OF THE NEW LAYOUT
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Gujarati:wght@400;600;700&display=swap');
    .a4-page {
        background-color: white; color: black; width: 210mm; min-height: 297mm;
        padding: 15mm 20mm; margin: 10px auto;
        font-family: 'Noto Sans Gujarati', sans-serif; font-size: 13px; line-height: 1.5;
        box-shadow: 0 0 15px rgba(0,0,0,0.5);
    }
    .input-line { border-bottom: 1px solid black; display: inline-block; padding: 0 5px; }
    table.budget-table { width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 12px; }
    table.budget-table th, table.budget-table td { border: 1px solid black; padding: 8px; vertical-align: middle; }
    table.budget-table th { text-align: center; font-weight: bold; }
    table.budget-table td { height: 35px; }
    .flex-row { display: flex; justify-content: space-between; align-items: flex-end; }
    .center { text-align: center; }
    .right { text-align: right; }
</style>
""", unsafe_allow_html=True)

page2_preview = f"""
<div class="a4-page">
    <div class="right" style="margin-bottom: 10px; font-size: 12px;">
    હિસાબી પત્રક નંબર <span class="input-line" style="width:100px;"></span>
    </div>
    <div class="center">
    <div style="font-size: 22px; font-weight: 700;">નવસારી કૃષિ વિશ્વવિધાલય</div>
    <div style="font-size: 16px; font-weight: 700; margin-top: 5px;">મુસાફરી ભથ્થા બીલ</div>
    </div>
    <div class="flex-row" style="margin-top: 20px;">
    <div style="width: 40%; font-weight: 700; font-size: 14px;">
    <div style="margin-bottom: 10px;">બીલ નંબર :</div>
    <div>તારીખ &nbsp;&nbsp;&nbsp;:</div>
    </div>
    <div style="width: 50%; text-align: right; font-size: 12px;">
    <div style="margin-bottom: 5px;">વાઉચર નં. <span class="input-line" style="width:180px;"></span></div>
    <div style="margin-bottom: 5px;">તારીખ <span class="input-line" style="width:180px;"></span></div>
    <div style="margin-bottom: 5px;">યુનિટ નંબર : <span class="input-line" style="width:165px;"></span></div>
    <div>કોડ નંબર : <span class="input-line" style="width:180px;"></span></div>
    </div>
    </div>
    <div class="center" style="margin-top: 20px; font-weight: 700; font-size: 12px;">
    આચાર્ય અને ડીનશ્રી, નં. મ. કૃષિ મહાવિદ્યાલય, નકૃયું, નવસારી ની કચેરીનું માહે: {voucher_month} નું
    </div>
    <div class="center" style="margin-top: 5px;">
    <span style="font-size: 16px; font-weight: 700; border-bottom: 2px solid black; padding-bottom: 2px;">મુસાફરી ભથ્થા બિલ</span>
    </div>
    <div style="margin-top: 20px; font-weight: 600; font-size: 12px;">
    યુનિટ/સબયુનિટ : આચાર્ય અને ડીનશ્રી, ન. મ. કૃષિ મહાવિદ્યાલય, નકૃયું, નવસારી
    </div>
    <div class="center" style="margin-top: 15px; font-weight: 600;">
    ખર્ચ માટેનું બજેટ સદર :-
    </div>
    <div class="center" style="margin-top: 5px; font-weight: 600;">
    યોજનાનું નામ :-
    </div>
    <div style="margin-top: 15px; font-size: 14px;">
    આથી રૂા. <span class="input-line" style="width: 150px; text-align:center; font-weight:bold;">{total_claim_val:.0f}</span> ના દાવો મંજુર કરી ગ્રાહય રાખવામાં આવે છે.
    </div>
    <div style="margin-top: 10px; border-top: 1px solid black; padding-top: 10px;">
    <div style="line-height: 2;">
    આ બીલમાં જણાવેલ રૂા <span class="input-line" style="width: 80px; text-align:center; font-weight:bold;">{total_claim_val:.0f}</span> (
    <span style="margin-left: 20px;">અંકે રૂપિયા <span style="font-weight:bold; font-size:14px;">{total_words_str} પુરા</span></span>
    <span style="float:right;">પૈસા)</span>
    <br>
    મંજુર કરવામાં આવે છે. અને તે રોકડા / ચેક નં. <span class="input-line" style="width: 150px;"></span> તા. <span class="input-line" style="width: 100px;"></span> થી ચુકવવામાં આવે છે.
    </div>
    </div>
    <hr style="border-top: 2px solid black; margin: 15px 0;">
    <table class="budget-table">
    <colgroup><col style="width: 70%;"><col style="width: 15%;"><col style="width: 15%;"></colgroup>
    <thead>
    <tr><th rowspan="2" style="text-align:left; border-bottom: none;"></th><th>રૂ.</th><th>પૈસા</th></tr>
    <tr><th style="height:20px;"></th><th style="height:20px;"></th></tr>
    </thead>
    <tbody>
    <tr><td>(૧) સને ૨૦૨૪-૨૫ માટે બજેટમાં મંજુર થયેલ રકમ</td><td></td><td></td></tr>
    <tr><td>(૨) આ બીલ સાથે થયેલ કુલ ખર્ચ</td><td></td><td></td></tr>
    <tr><td>(૩) ખર્ચ માટે બાકી રહેતી રકમ</td><td></td><td></td></tr>
    </tbody>
    </table>
</div>
"""

page3_preview = f"""
<div class="a4-page">
    <div style="font-size: 11px;">
    <strong>નોંધ :-</strong>
    <ol style="margin-top: 5px; padding-left: 20px; margin-bottom: 10px;">
    <li>કોલમ નં. ૭ માં મુસાફરી પ્રકાર રેલ્વે/એસ.ટી./હવાઈ/સ્ટીમર/ભાડાનું યુનિવર્સિટી કે સરકારી કે પોતાનું વાહન ઈત્યાદી મારફત કરેલ મુસાફરીની સ્પષ્ટ નોંધ આપવી.</li>
    <li>કોલમ નં. ૧૧ થી ૧૩ માઈલેજ મેળવતા અધિકારીઓ કે સભ્યોએ ભરવી.</li>
    </ol>
    </div>
    <div style="margin-top: 10px;">
    <strong>યુનિવર્સિટી કર્મચારીએ આપવાનું પ્રમાણપત્ર</strong>
    <ol style="margin-top: 5px; padding-left: 20px;">
    <li>આથી પ્રમાણપત્ર આપવામાં આવે છે કે, આ બીલમાં આકારેલ રકમ બીજા કોઈ બીલમાં આકારેલ નથી.</li>
    <li>આથી પ્રમાણીત કરવામાં આવે છે કે સદર મુસાફરી ભથ્થા બીલમાં દર્શાવેલ હકીકત સાચી છે...</li>
    <li>આથી પ્રમાણપત્ર આપવામાં આવે છે કે બીલમાં દર્શાવેલ પ્રવાસ માટે મેં આ અગાઉ પેશગી લીધેલ નથી...</li>
    </ol>
    </div>
    <div class="flex-row" style="margin-top: 20px;">
    <div></div>
    <div style="text-align: center;">
    <strong>({emp_name})</strong><br>
    {designation}<br>
    કર્મચારીની સહી નામ અને હોદ્દો
    </div>
    </div>
    <div style="margin-top: 20px; font-weight: 700; font-size: 13px; border-bottom: 2px solid black; padding-bottom: 2px;">
    કર્મચારી / અધિકારી / સભ્યશ્રીએ નીચેની વિગત ભરવી.
    </div>
    <div style="display: flex; border-bottom: 2px solid black;">
    <div style="width: 50%; border-right: 2px solid black; padding: 10px 10px 10px 0; font-weight: 600;">
    <div style="display: flex; justify-content: space-between; margin-bottom: 5px;">
    <span>બીલની કુલ રકમ</span><span style="font-weight: 700; font-size: 14px;">= &nbsp;&nbsp; {total_claim_val:.0f}</span>
    </div>
    <div style="display: flex; justify-content: space-between; margin-bottom: 15px;">
    <span>ચૂકવવા પાત્ર ચોખ્ખી રકમ</span><span style="font-weight: 700; font-size: 14px;">= &nbsp;&nbsp; {total_claim_val:.0f}</span>
    </div>
    </div>
    <div style="width: 50%; padding: 10px 0 10px 10px; position: relative; font-weight: 600;">
    <div style="margin-bottom: 5px;">બીલની રકમ રૂ. <span style="font-weight: 700; border-bottom: 1px solid black; font-size: 14px;">{total_claim_val:.0f}</span></div>
    <div style="margin-bottom: 15px;">અંકે રૂપિયા <span style="font-weight: 700; border-bottom: 1px solid black; font-size: 14px;">{total_words_str} પુરા</span></div>
    <div style="margin-bottom: 20px;">મને મળ્યા છે.</div>
    <div style="text-align: center; position: absolute; bottom: 10px; right: 10px;">
    <div style="font-weight: 700;">({emp_name})</div>
    </div>
    </div>
    </div>
</div>
"""

# --- UI PREVIEW RENDER ---
st.markdown("### 📄 Layout Preview")

col1, col2 = st.columns(2)
with col1:
    st.info("Page 1 (Hisabi Patrak)")
    st.markdown(page2_preview, unsafe_allow_html=True)

with col2:
    st.info("Page 4 (Notes & Receipt)")
    st.markdown(page3_preview, unsafe_allow_html=True)

# Button
if st.button("📄 Generate & Download Final Word File", type="primary"):
    if not df_ta.empty:
        # Generate
        doc = create_split_layout_doc(df_ta)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="⬇️ Download Final_TA_Bill_Split.docx",
            data=buffer,
            file_name="Final_TA_Bill_Split.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.balloons()
