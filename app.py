import streamlit as st
import google.generativeai as genai
import pandas as pd
import json
import os
from PIL import Image
import PyPDF2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# --- CONFIGURATION ---
AI_MODEL_NAME = "gemini-3-flash-preview" 
ST_DATA_DIR = "data_store"
ST_RULES_DIR = "rules_store"
os.makedirs(ST_DATA_DIR, exist_ok=True)
os.makedirs(ST_RULES_DIR, exist_ok=True)
HISTORY_FILE = os.path.join(ST_DATA_DIR, "claim_history.json")

st.set_page_config(page_title="NAU TA/DA Claim Assistant", layout="wide", page_icon="🚜")

# --- SESSION STATE INITIALIZATION ---
def init_session():
    # Tour Data: Detailed leg-by-leg travel
    if "tour_data" not in st.session_state:
        st.session_state.tour_data = pd.DataFrame(columns=[
            "Dep_Date", "Dep_Time", "Dep_Place", 
            "Arr_Date", "Arr_Time", "Arr_Place", 
            "Mode", "Vehicle_Details", "Distance_KM", 
            "Ticket_Amount", "Enquiry_Fare", "Remark"
        ])
    
    # Stay Data: Hotel/Guest House details
    if "stay_data" not in st.session_state:
        st.session_state.stay_data = pd.DataFrame(columns=[
            "CheckIn_Date", "CheckIn_Time", 
            "CheckOut_Date", "CheckOut_Time", 
            "City", "Stay_Type", "Bill_Amount", "Remark"
        ])
    
    if "salary_info" not in st.session_state:
        st.session_state.salary_info = {"Basic Pay": 0, "Pay Level": "Level 12", "Designation": "Associate Professor"}

init_session()

# --- HELPER FUNCTIONS ---
def get_gemini_response(prompt, content_parts, api_key):
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(AI_MODEL_NAME) 
        response = model.generate_content([prompt, *content_parts])
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

def extract_text_from_pdf(uploaded_file):
    reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def calculate_da_rate(pay_level, city_class, stay_type):
    """
    Returns (Rate, Rule_Text) based on Revised Provisions 2024.
    Rates are hardcoded based on the PDF rules provided.
    """
    # Parse Pay Level (Simplified logic)
    try:
        level_num = int(''.join(filter(str.isdigit, str(pay_level))))
    except:
        level_num = 10 # Default fallback
    
    # Logic for Level 12 and above
    if level_num >= 12:
        if stay_type == "Hotel":
            if city_class == "X": return 2600, "Level 12+, X City, Hotel"
            if city_class == "Y": return 2100, "Level 12+, Y City, Hotel"
            return 1300, "Level 12+, Z City, Hotel"
        else: # Regular/Guest House
            if city_class == "X": return 1000, "Level 12+, X City, Ordinary"
            if city_class == "Y": return 800, "Level 12+, Y City, Ordinary"
            return 500, "Level 12+, Z City, Ordinary"

    # Logic for Level 6 to 11
    elif 6 <= level_num <= 11:
        if stay_type == "Hotel":
            if city_class == "X": return 2000, "Level 6-11, X City, Hotel"
            if city_class == "Y": return 1600, "Level 6-11, Y City, Hotel"
            return 900, "Level 6-11, Z City, Hotel"
        else:
            if city_class == "X": return 900, "Level 6-11, X City, Ordinary"
            if city_class == "Y": return 700, "Level 6-11, Y City, Ordinary"
            return 400, "Level 6-11, Z City, Ordinary"
            
    # Logic for Level 5 and below
    else:
        if stay_type == "Hotel":
            if city_class == "X": return 900, "Level 5-, X City, Hotel"
            if city_class == "Y": return 800, "Level 5-, Y City, Hotel"
            return 500, "Level 5-, Z City, Hotel"
        else:
            if city_class == "X": return 700, "Level 5-, X City, Ordinary"
            if city_class == "Y": return 500, "Level 5-, Y City, Ordinary"
            return 300, "Level 5-, Z City, Ordinary"

def get_city_class(city_name):
    # Simplified Dictionary based on Circular 1734
    x_cities = ["Ahmedabad", "Bengaluru", "Chennai", "Delhi", "Hyderabad", "Kolkata", "Mumbai", "Pune", "Gandhinagar"]
    y_cities = ["Vadodara", "Surat", "Rajkot", "Jamnagar", "Bhavnagar", "Anand", "Navsari"] # Added common GJ cities
    
    city_norm = str(city_name).strip().title()
    for x in x_cities:
        if x in city_norm: return "X"
    for y in y_cities:
        if y in city_norm: return "Y"
    return "Z"

# --- MAIN UI ---
st.title("🚜 NAU TA/DA Claim Assistant")

# Sidebar
with st.sidebar:
    st.header("Settings")
    api_key = st.text_input("Gemini API Key", type="password")
    st.divider()
    st.markdown("### 📋 User Profile")
    st.session_state.salary_info["Designation"] = st.text_input("Designation", st.session_state.salary_info.get("Designation"))
    st.session_state.salary_info["Pay Level"] = st.selectbox("Pay Level", ["Level 14", "Level 13A", "Level 13", "Level 12", "Level 11", "Level 10", "Level 9", "Level 8", "Level 7"], index=3)
    st.session_state.salary_info["Basic Pay"] = st.number_input("Basic Pay", value=st.session_state.salary_info.get("Basic Pay", 0))

# Tabs
tab1, tab2, tab3, tab4 = st.tabs(["1. Upload & Extract", "2. Edit Tour & Stay", "3. Review Calculations", "4. Export Doc"])

# --- TAB 1: UPLOADS ---
with tab1:
    col1, col2 = st.columns(2)
    with col1:
        st.info("Upload Tour Diary / Tickets")
        files = st.file_uploader("Upload Files (PDF/Images)", accept_multiple_files=True)
    
    if files and st.button("🚀 Auto-Extract Journey Data"):
        if not api_key:
            st.error("Please enter API Key.")
        else:
            with st.spinner("Extracting journey details..."):
                content = ["Extract all travel legs and stay details."]
                for f in files:
                    if f.type == "application/pdf":
                        content.append(extract_text_from_pdf(f))
                    else:
                        content.append(Image.open(f))
                
                prompt = """
                Extract journey details into a JSON object with two lists: 'tour' and 'stay'.
                
                For 'tour' (Travel Legs):
                [{"Dep_Date": "YYYY-MM-DD", "Dep_Time": "HH:MM", "Dep_Place": "City", 
                  "Arr_Date": "YYYY-MM-DD", "Arr_Time": "HH:MM", "Arr_Place": "City",
                  "Mode": "Bus/Rail/Air/Private/Taxi", "Distance_KM": 0, "Ticket_Amount": 0, "Enquiry_Fare": 0}]
                
                For 'stay' (Accommodation):
                [{"CheckIn_Date": "YYYY-MM-DD", "CheckIn_Time": "HH:MM", 
                  "CheckOut_Date": "YYYY-MM-DD", "CheckOut_Time": "HH:MM",
                  "City": "City Name", "Stay_Type": "Hotel/Guest House", "Bill_Amount": 0}]
                """
                
                res = get_gemini_response(prompt, content, api_key)
                try:
                    clean = res.replace("```json", "").replace("```", "")
                    data = json.loads(clean)
                    
                    if "tour" in data:
                        st.session_state.tour_data = pd.DataFrame(data["tour"])
                    if "stay" in data:
                        st.session_state.stay_data = pd.DataFrame(data["stay"])
                    st.success("Extraction Complete! Go to 'Edit' tab.")
                except Exception as e:
                    st.error(f"Extraction failed: {e}")

# --- TAB 2: EDIT DATA ---
with tab2:
    st.markdown("### 🗓️ Chronological Tour Data")
    st.caption("Please ensure Date format is YYYY-MM-DD and Time is HH:MM")
    
    # 1. TOUR DATA EDITOR
    st.subheader("A. Journey Details (TA)")
    
    mode_options = [
        "Government Bus (GSRTC)", "Railway", "Air", 
        "Road - Private Vehicle (Petrol)", "Road - Private Vehicle (Diesel)", 
        "Taxi/Auto", "Official Vehicle"
    ]
    
    column_config_tour = {
        "Dep_Date": st.column_config.DateColumn("Dep. Date"),
        "Dep_Time": st.column_config.TimeColumn("Dep. Time"),
        "Arr_Date": st.column_config.DateColumn("Arr. Date"),
        "Arr_Time": st.column_config.TimeColumn("Arr. Time"),
        "Mode": st.column_config.SelectboxColumn("Mode", options=mode_options, required=True),
        "Vehicle_Details": st.column_config.TextColumn("Vehicle Details", help="e.g. Car No. GJ-05..."),
        "Ticket_Amount": st.column_config.NumberColumn("Ticket Paid (Rs)"),
        "Enquiry_Fare": st.column_config.NumberColumn("Enquiry Fare (Rs)", help="REQUIRED for Private Vehicle claims"),
    }
    
    edited_tour = st.data_editor(
        st.session_state.tour_data,
        column_config=column_config_tour,
        num_rows="dynamic",
        use_container_width=True,
        key="tour_editor"
    )
    st.session_state.tour_data = edited_tour

    st.divider()

    # 2. STAY DATA EDITOR
    st.subheader("B. Stay / Accommodation Details (DA)")
    
    column_config_stay = {
        "CheckIn_Date": st.column_config.DateColumn("Check-In"),
        "CheckOut_Date": st.column_config.DateColumn("Check-Out"),
        "Stay_Type": st.column_config.SelectboxColumn("Type", options=["Hotel", "Guest House", "Own Arrangement"]),
        "City": st.column_config.TextColumn("City", help="Determines X, Y, Z Class"),
        "Bill_Amount": st.column_config.NumberColumn("Bill Amount")
    }
    
    edited_stay = st.data_editor(
        st.session_state.stay_data,
        column_config=column_config_stay,
        num_rows="dynamic",
        use_container_width=True,
        key="stay_editor"
    )
    st.session_state.stay_data = edited_stay

# --- TAB 3: CALCULATIONS ---
with tab3:
    st.header("🧮 Claims Calculation")
    
    if st.button("Calculate TA & DA"):
        # 1. TA CALCULATION
        st.subheader("1. Travelling Allowance (TA)")
        ta_claim_data = []
        total_ta = 0
        
        for idx, row in st.session_state.tour_data.iterrows():
            mode = str(row.get("Mode", ""))
            ticket = float(row.get("Ticket_Amount", 0) or 0)
            enquiry = float(row.get("Enquiry_Fare", 0) or 0)
            claimed = 0
            rule = ""
            
            if "Private Vehicle" in mode:
                claimed = enquiry
                rule = "Private Vehicle: Restricted to Official Enquiry Fare (No Mileage)"
            elif "Official Vehicle" in mode:
                claimed = 0
                rule = "Official Vehicle: No Claim Admissible"
            else:
                claimed = ticket
                rule = "Public Transport: Actual Ticket Amount"
            
            total_ta += claimed
            ta_claim_data.append({
                "Date": row.get("Dep_Date"),
                "From": row.get("Dep_Place"),
                "To": row.get("Arr_Place"),
                "Mode": mode,
                "Claimed": claimed,
                "Rule Applied": rule
            })
            
        st.dataframe(pd.DataFrame(ta_claim_data))
        st.metric("Total TA Claim", f"₹ {total_ta}")
        
        # 2. DA CALCULATION
        st.subheader("2. Daily Allowance (DA)")
        da_claim_data = []
        total_da = 0
        
        # Process Stays
        for idx, row in st.session_state.stay_data.iterrows():
            city = row.get("City", "Gandhinagar")
            stay_type = row.get("Stay_Type", "Guest House")
            
            # Determine Class and Rate
            c_class = get_city_class(city)
            rate, rule_desc = calculate_da_rate(st.session_state.salary_info["Pay Level"], c_class, stay_type)
            
            # Simple day diff calculation (Enhance logic for partial days if needed)
            try:
                d1 = pd.to_datetime(row["CheckIn_Date"])
                d2 = pd.to_datetime(row["CheckOut_Date"])
                days = (d2 - d1).days
                if days < 1: days = 1 # Minimum 1 day if dates are same
            except:
                days = 0
                
            amount = rate * days
            total_da += amount
            
            da_claim_data.append({
                "City": city,
                "Class": c_class,
                "Stay Type": stay_type,
                "Rate": rate,
                "Days": days,
                "Total": amount,
                "Rule": rule_desc
            })
            
        st.dataframe(pd.DataFrame(da_claim_data))
        st.metric("Total DA Claim", f"₹ {total_da}")
        
        st.success(f"GRAND TOTAL: ₹ {total_ta + total_da}")
        
        # Save calculated totals for export
        st.session_state.calc_totals = {"TA": total_ta, "DA": total_da, "Grand": total_ta + total_da}
        st.session_state.ta_details = ta_claim_data
        st.session_state.da_details = da_claim_data

# --- TAB 4: EXPORT ---
with tab4:
    st.header("📄 Export Tour Diary")
    
    if st.button("Generate Word Document (.docx)"):
        doc = Document()
        
        # Title
        title = doc.add_paragraph("NAU TOUR DIARY & CLAIM FORM")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = doc.styles['Title']
        
        # Personal Info
        # Using a get() with a default to avoid errors if the key is missing
        doc.add_paragraph(f"Name: {st.session_state.get('user_name', 'Vaibhav Kumar Kanubhai Chaudhari')}")
        doc.add_paragraph(f"Designation: {st.session_state.salary_info.get('Designation')}")
        doc.add_paragraph(f"Pay Level: {st.session_state.salary_info.get('Pay Level')}")
        doc.add_paragraph(f"Headquarters: Navsari Agricultural University")
        
        # 1. Journey Table
        doc.add_heading('1. Details of Journey (TA)', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date/Time'
        hdr_cells[1].text = 'From'
        hdr_cells[2].text = 'To'
        hdr_cells[3].text = 'Mode'
        hdr_cells[4].text = 'Rule Applied'
        hdr_cells[5].text = 'Amount (Rs)'
        
        # Sort tour data by date before printing
        sorted_tour = st.session_state.tour_data.sort_values(by="Dep_Date")
        
        for idx, row in sorted_tour.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = f"{row.get('Dep_Date')}\n{row.get('Dep_Time')}"
            row_cells[1].text = str(row.get('Dep_Place'))
            row_cells[2].text = str(row.get('Arr_Place'))
            
            mode_txt = str(row.get('Mode'))
            if "Private" in mode_txt:
                mode_txt += f"\n({row.get('Vehicle_Details', '')})"
            row_cells[3].text = mode_txt
            
            # Logic for amount and rule
            amt = row.get('Ticket_Amount', 0)
            rule_note = "Actual Ticket"
            if "Private" in mode_txt:
                amt = row.get('Enquiry_Fare', 0)
                rule_note = "Restricted to Fare Enquiry"
            
            row_cells[4].text = rule_note
            row_cells[5].text = str(amt)
            
        # 2. Stay Table
        doc.add_heading('2. Daily Allowance (DA)', level=2)
        da_table = doc.add_table(rows=1, cols=5)
        da_table.style = 'Table Grid'
        da_hdr = da_table.rows[0].cells
        da_hdr[0].text = 'City'
        da_hdr[1].text = 'Class (X/Y/Z)'
        da_hdr[2].text = 'Stay Type'
        da_hdr[3].text = 'Rate'
        da_hdr[4].text = 'Days'
        
        if "da_details" in st.session_state:
            for item in st.session_state.da_details:
                row_cells = da_table.add_row().cells
                row_cells[0].text = str(item['City'])
                row_cells[1].text = str(item['Class'])
                row_cells[2].text = str(item['Stay Type'])
                row_cells[3].text = str(item['Rate'])
                row_cells[4].text = str(item['Days'])

        # Totals
        doc.add_paragraph(f"\nTotal TA Claim: Rs. {st.session_state.get('calc_totals', {}).get('TA', 0)}")
        doc.add_paragraph(f"Total DA Claim: Rs. {st.session_state.get('calc_totals', {}).get('DA', 0)}")
        grand_p = doc.add_paragraph(f"GRAND TOTAL: Rs. {st.session_state.get('calc_totals', {}).get('Grand', 0)}")
        grand_p.runs[0].bold = True
        
        # Definitions Footer - FULL TEXT INSERTION
        doc.add_page_break()
        doc.add_heading('Definitions & Rules Applied', level=3)
        
        definitions = [
            ("Mode of Transport (Railway / Bus / Flight)", "Means the actual means of conveyance used by the employee for performing the journey on official tour, such as Railway, Public Bus, or Air, as recorded in the tour diary and supported by valid travel documents."),
            ("Class of Travel (I / II / III)", "Means the category of travel entitlement applicable to a University employee for performing a journey on official tour, such as First Class, Second Class, or Third Class, as per the employee’s pay level, designation, and eligibility prescribed under the Travelling Allowance Rules."),
            ("Ticket Price / Rate (Rs.)", "Means the actual fare paid in rupees for travel by the entitled mode and class for performing a journey on official tour, as supported by the original travel ticket, receipt, or authorised proof of fare. Where the original ticket or receipt is not available, the fare determined through an official fare enquiry from the Railway / State Transport Bus / Airline for the same date, route, and eligible class of travel may be considered for the purpose of Travelling Allowance calculation, subject to admissibility under the rules and certification by the competent authority."),
            ("Road Travel by Other Vehicle", "Means travel performed on official tour by a mode of road transport other than Railway or Air, such as State Transport Bus, Metro, Auto-rickshaw, Taxi, or other public conveyance, used for the journey between places connected by road. In cases where travel is performed by a motor vehicle powered by diesel or petrol, the type of vehicle and fuel used shall be clearly indicated in the tour diary. However, use of a private vehicle is not ordinarily admissible for reimbursement in the University, and in such cases, Travelling Allowance shall be restricted to the fare determined through an official fare enquiry of the eligible public transport for the same route and distance. Reimbursement for road travel shall be regulated on the basis of fare enquiry or admissible public conveyance rates, and not on the basis of private vehicle ownership, even though provisions for road mileage exist under the rules. Travel by auto-rickshaw, metro rail, city bus, or other recognised public transport is admissible for calculation of Travelling Allowance, subject to necessity, reasonableness, and certification in the tour diary."),
            ("Days of Daily Allowance receivable", "Means the number of days for which Daily Allowance (DA) is admissible to a University employee while on tour, determined on the basis of the total duration of absence from headquarters, including journey time, and calculated in accordance with the minimum time limits prescribed under the Travelling Allowance Rules."),
            ("Daily Allowance Rate (Rs.)", "Means the monetary rate of Daily Allowance prescribed in rupees and applicable to a University employee for a day of tour, determined on the basis of the employee’s pay level/grade and the classification of the city or place of halt, as notified under the Travelling Allowance provisions.")
        ]
        
        for title, text in definitions:
            p = doc.add_paragraph()
            p.add_run(title + ": ").bold = True
            p.add_run(text)
            # Add a small spacing after each definition
            p.paragraph_format.space_after = Pt(12)

        # Save to buffer
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="Download Tour Diary (.docx)",
            data=buffer,
            file_name="NAU_Tour_Diary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

